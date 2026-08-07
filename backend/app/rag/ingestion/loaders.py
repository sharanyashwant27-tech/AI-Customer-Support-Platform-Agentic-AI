"""Document loaders for PDF, DOCX, HTML, Markdown, and plain text (emails)."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from app.rag.cleaning import clean_document_text
from app.rag.sources import infer_knowledge_source


@dataclass
class LoadedDocument:
    title: str
    content: str
    source: str
    file_type: str
    metadata: dict[str, Any] = field(default_factory=dict)


def _with_source(doc: LoadedDocument) -> LoadedDocument:
    doc.metadata.setdefault(
        "knowledge_source",
        infer_knowledge_source(filename=doc.source, file_type=doc.file_type),
    )
    return doc


def load_markdown(content: str | bytes, *, title: str, source: str) -> LoadedDocument:
    text = content.decode("utf-8", errors="ignore") if isinstance(content, bytes) else content
    return _with_source(
        LoadedDocument(
            title=title,
            content=clean_document_text(text),
            source=source,
            file_type="markdown",
            metadata={"format": "md"},
        )
    )


def load_html(content: str | bytes, *, title: str, source: str) -> LoadedDocument:
    text = content.decode("utf-8", errors="ignore") if isinstance(content, bytes) else content
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(text, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        extracted = soup.get_text(separator="\n")
    except Exception:
        extracted = re.sub(r"<[^>]+>", " ", text)
    return _with_source(
        LoadedDocument(
            title=title,
            content=clean_document_text(extracted),
            source=source,
            file_type="html",
            metadata={"format": "html"},
        )
    )


def load_pdf(data: bytes, *, title: str, source: str) -> LoadedDocument:
    try:
        from pypdf import PdfReader
        from io import BytesIO

        reader = PdfReader(BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n\n".join(pages)
    except Exception as exc:
        text = f"[PDF parse unavailable: {exc}]"
    return _with_source(
        LoadedDocument(
            title=title,
            content=clean_document_text(text),
            source=source,
            file_type="pdf",
            metadata={"format": "pdf", "knowledge_source": "pdfs"},
        )
    )


def load_docx(data: bytes, *, title: str, source: str) -> LoadedDocument:
    try:
        from docx import Document
        from io import BytesIO

        doc = Document(BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
    except Exception as exc:
        text = f"[DOCX parse unavailable: {exc}]"
    return _with_source(
        LoadedDocument(
            title=title,
            content=clean_document_text(text),
            source=source,
            file_type="docx",
            metadata={"format": "docx"},
        )
    )


def load_email(content: str | bytes, *, title: str, source: str) -> LoadedDocument:
    text = content.decode("utf-8", errors="ignore") if isinstance(content, bytes) else content
    lines = []
    for line in text.splitlines():
        if re.match(r"^(From|To|Cc|Bcc|Date|Subject|MIME-Version|Content-Type):", line, re.I):
            if line.lower().startswith("subject:"):
                lines.append(line)
            continue
        lines.append(line)
    return _with_source(
        LoadedDocument(
            title=title,
            content=clean_document_text("\n".join(lines)),
            source=source,
            file_type="email",
            metadata={"format": "email", "knowledge_source": "emails"},
        )
    )


def load_bytes(
    data: bytes,
    *,
    filename: str,
    content_type: str | None = None,
) -> LoadedDocument:
    name = filename or "document"
    lower = name.lower()
    source = name
    if lower.endswith(".pdf") or (content_type and "pdf" in content_type):
        return load_pdf(data, title=name, source=source)
    if lower.endswith(".docx") or (
        content_type and "wordprocessingml" in content_type
    ):
        return load_docx(data, title=name, source=source)
    if lower.endswith((".html", ".htm")) or (content_type and "html" in content_type):
        return load_html(data, title=name, source=source)
    if lower.endswith((".eml", ".email")) or (
        content_type and "message/rfc822" in content_type
    ):
        return load_email(data, title=name, source=source)
    if lower.endswith((".md", ".markdown")) or (
        content_type and "markdown" in (content_type or "")
    ):
        return load_markdown(data, title=name, source=source)
    return load_markdown(data, title=name, source=source)


def load_path(path: Path) -> LoadedDocument:
    data = path.read_bytes()
    return load_bytes(data, filename=path.name)
